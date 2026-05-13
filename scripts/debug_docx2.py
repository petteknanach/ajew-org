#!/usr/bin/env python3
"""Debug: show paragraph counts for all docx files."""
from docx import Document
import os

DOCX_DIR = '/mnt/c/Users/Pettek/Documents/Claude Desktop projects/Finished/temp folder/Likutay Halachos'

for df in sorted(os.listdir(DOCX_DIR)):
    if not df.endswith('.docx'): continue
    doc = Document(os.path.join(DOCX_DIR, df))
    total = len(doc.paragraphs)
    he_count = 0
    en_count = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if len(t) < 15: continue
        has_he = any('\u05D0' <= c <= '\u05EA' for c in t)
        if has_he:
            he_count += 1
        else:
            en_count += 1
    print(f"{df:55s}: total={total:4d} HE={he_count:4d} EN={en_count:4d} (content paras >= 15 chars)")