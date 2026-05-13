#!/usr/bin/env python3
"""
Analyze what's in each JSON segment vs what's in the docx.
Focus on understanding the actual mapping needed.
"""
from docx import Document
import json, os, re

DOCX_DIR = '/mnt/c/Users/Pettek/Documents/Claude Desktop projects/Finished/temp folder/Likutay Halachos'
READER_DIR = '/root/ajew-org/public/reader/likutay-halachos'

# What does the JSON contain?
print("=== JSON segment analysis ===\n")
for part in ['part-1', 'part-2']:
    part_path = os.path.join(READER_DIR, part)
    jsfiles = sorted([f for f in os.listdir(part_path)
                      if f.endswith('.json') and f != 'index.json'])

    for jf in jsfiles[:3]:  # First 3 files
        data = json.load(open(os.path.join(part_path, jf)))
        print(f"{part}/{jf}: {len(data['segments'])} segments")

        for i, seg in enumerate(data['segments'][:3]):
            he = seg.get('he','').strip()
            en = seg.get('en','').strip()
            he_chars = sum(1 for c in he if '\u05D0' <= c <= '\u05EA')
            he_ratio = he_chars / len(he) if he else 0

            print(f"  Seg {i+1}:")
            print(f"    HE ({len(he)} chars, {he_ratio:.0%} Hebrew): {he[:60]}")
            print(f"    EN ({len(en)} chars): {en[:60] if en else '(empty)'}")
        print()

# What does the docx actually contain?
print("\n=== Docx content analysis ===\n")
doc = Document(os.path.join(DOCX_DIR, 'Volume_01_OC1_English.docx'))

# Find paragraphs around the content start
content_start = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    # Look for "The Laws of Rising in the Morning" which is the first actual law
    if 'Rising in the Morning' in t:
        content_start = i
        break

print(f"First law paragraph at: P{content_start}")

# Show the structure around that point
for i in range(max(0, content_start - 5), min(len(doc.paragraphs), content_start + 30)):
    t = doc.paragraphs[i].text.strip()
    if not t: continue
    has_he = any('\u05D0' <= c <= '\u05EA' for c in t)
    print(f"P{i:4d} [{'HE' if has_he else 'EN'}] | {t[:100]}")