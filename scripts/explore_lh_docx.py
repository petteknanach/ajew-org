#!/usr/bin/env python3
"""Explore structure of Likutay Halachos docx files to understand EN-HE pairing."""
from docx import Document
import re

# Check first volume in detail
doc = Document('/mnt/c/Users/Pettek/Documents/Claude Desktop projects/Finished/temp folder/Likutay Halachos/Volume_01_OC1_English.docx')

print("LH docx structure analysis:")
print("=" * 80)
print(f'Total paragraphs: {len(doc.paragraphs)}')
print()

for i, p in enumerate(doc.paragraphs[:80]):
    text = p.text.strip()
    if not text:
        continue
    # Check paragraph style
    style = p.style.name if p.style else 'None'
    # Check if it has runs
    runs_text = ''.join(r.text for r in p.runs)
    hebrew = any(ord(c) > 127 for c in runs_text)

    print(f'P{i}: style=[{style}] len={len(text)} hebrew={hebrew}')
    print(f'    {text[:200]}')
    print()