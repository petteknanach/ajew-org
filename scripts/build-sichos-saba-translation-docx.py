#!/usr/bin/env python3
"""Build the tape/side-ordered Sichos Saba English translation DOCX.

Strict mode refuses to create a document until every available canonical source
side has reviewed English. --allow-incomplete is for internal layout QA only.
"""
from __future__ import annotations
import argparse, json
from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
BOOK = ROOT / "public/reader/saba-tape-transcripts"
DEFAULT_OUT = Path("/mnt/c/Users/Pettek/Downloads/Sichos_Saba_Complete_English_Translation.docx")
SIDE_NAME = {"a": "Side A (Aleph)", "b": "Side B (Bais)"}


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def legacy_paragraphs(existing: dict) -> list[str]:
    chapter = BOOK / f"chapter-{existing['chapter']}.json"
    data = json.loads(chapter.read_text(encoding="utf-8-sig"))
    return [s.get("en", "").strip() for s in data.get("segments", []) if s.get("en", "").strip()]


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--output", type=Path, default=DEFAULT_OUT)
    ap.add_argument("--allow-incomplete", action="store_true")
    args = ap.parse_args()
    manifest = json.loads((BOOK / "translation-manifest.json").read_text(encoding="utf-8"))

    assembled: list[tuple[dict, list[str], str]] = []
    incomplete: list[str] = []
    for item in manifest["sides"]:
        data = json.loads((BOOK / item["file"]).read_text(encoding="utf-8"))
        label = f"Tape {item['tape']} {SIDE_NAME[item['side']]}"
        status = data.get("translationStatus")
        paragraphs: list[str] = []
        if status == "verified":
            paragraphs = [s["en"].strip() for s in data.get("segments", []) if s.get("en", "").strip()]
            if len(paragraphs) != len(data.get("segments", [])):
                incomplete.append(label)
        elif status == "existing_verified" and data.get("existingEnglish"):
            paragraphs = legacy_paragraphs(data["existingEnglish"])
        elif data.get("sourceStatus") == "missing":
            paragraphs = ["[The supplied transcript explicitly states that this tape side is missing. No wording has been invented.]"]
        else:
            incomplete.append(label)
            if args.allow_incomplete:
                if data.get("existingEnglish"):
                    paragraphs = legacy_paragraphs(data["existingEnglish"])
                    status = "existing_review_pending"
                else:
                    paragraphs = ["[English translation pending verified completion.]"]
        assembled.append((data, paragraphs, status or "not_started"))

    if incomplete and not args.allow_incomplete:
        raise SystemExit(f"REFUSED: {len(incomplete)} tape sides lack reviewed complete English; first: {', '.join(incomplete[:10])}")

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.bottom_margin = Inches(0.75)
    sec.left_margin = sec.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"; normal.font.size = Pt(11.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for name, size in (("Title", 24), ("Heading 1", 18), ("Heading 2", 14)):
        style = doc.styles[name]
        style.font.name = "Times New Roman"; style.font.size = Pt(size)
    add_page_number(sec.footer.paragraphs[0])

    title = doc.add_heading("Sichos Saba", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Complete English Translation — Organized by Tape and Tape Side")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Teachings and conversations of Saba Yisroel, Rabbi Yisroel Dov Odesser")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    status_text = "Complete reviewed edition" if not incomplete else f"Internal progress proof — {len(incomplete)} tape sides still require reviewed English"
    p = doc.add_paragraph(status_text); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Source note: The supplied transcript warns that it contains transcription errors and that unclear readings should be checked against the recordings. Explicitly missing and partial source material is identified rather than reconstructed or summarized.")

    last_tape = None
    for data, paragraphs, status in assembled:
        tape = data["tapeNumber"]
        if tape != last_tape:
            doc.add_page_break()
            doc.add_heading(f"Tape {tape}", level=1)
            last_tape = tape
        doc.add_heading(SIDE_NAME[data["side"]], level=2)
        if data.get("sourceStatus") == "partial":
            doc.add_paragraph("[Source notice: the supplied transcript states that much of this side is missing or inaudible. Only available wording is translated.]")
        if status == "existing_review_pending" and args.allow_incomplete:
            doc.add_paragraph("[Existing English tape-side translation included for layout review; canonical alignment review pending.]")
        for block in paragraphs:
            for para_text in block.split("\n\n"):
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    print(json.dumps({
        "output": str(args.output),
        "tapes": 117,
        "sides": len(assembled),
        "incompleteSides": len(incomplete),
        "strict": not args.allow_incomplete,
        "bytes": args.output.stat().st_size,
    }, indent=2))

if __name__ == "__main__":
    main()
