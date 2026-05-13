#!/usr/bin/env python3
"""
Fix LH pairing v15 - Clean approach.

For each part, extract EN paragraphs from docx volumes that belong to that part.
Then assign them to JSON content segments by position.
"""
from docx import Document
import json, os, re

DOCX_DIR = '/mnt/c/Users/Pettek/Documents/Claude Desktop projects/Finished/temp folder/Likutay Halachos'
READER_DIR = '/root/ajew-org/public/reader/likutay-halachos'

def is_hebrew_char(c): return '\u05D0' <= c <= '\u05EA'
def has_hebrew(t): return any(is_hebrew_char(c) for c in t)

def is_content_en(text):
    """Check if this is an actual English translation paragraph (not metadata)."""
    t = text.strip()
    if len(t) < 20: return False
    # Must have mostly ASCII letters
    alpha = [c for c in t if c.isalpha()]
    if not alpha: return False
    ascii_ratio = sum(1 for c in alpha if c.isascii()) / len(alpha)
    if ascii_ratio < 0.6: return False

    # Skip known metadata patterns
    meta = ['hilchos ', 'na nach', 'siman ', 'seif ', 'osio ',
            'volume ', 'introduction', 'likutay', 'a collection',
            'the laws ', 'oc ', 'yd ', 'eh ', 'cm ', 'like all',
            'naanach', 'segment', 'one stop', 'arranged by',
            'copyright', 'rough draft', 'no copyright',
            'free for all', 'books of rabbi', 'na na',
            'table of contents', 'right-click', 'update field',
            'each paragraph in this', 'cross-reference',
            'bas sheva', 'yisroel dov', 'student of rabbi',
            'note on paragraph', 'rabbi nachman']
    t_lower = t.lower()
    if any(m in t_lower for m in meta):
        return False
    return True

def extract_en_for_volume(vol_num):
    """Extract EN content paragraphs from a specific volume docx."""
    # Find the docx file
    for df in sorted(os.listdir(DOCX_DIR)):
        if not df.endswith('.docx'): continue
        # Match volume number
        vol_match = re.search(r'Volume_(\d+)', df)
        if not vol_match or int(vol_match.group(1)) != vol_num:
            continue

        path = os.path.join(DOCX_DIR, df)
        doc = Document(path)

        # Find content start
        content_start = 0
        for i in range(len(doc.paragraphs)):
            if is_content_en(doc.paragraphs[i].text):
                content_start = i
                break

        # Extract all content EN paragraphs after content start
        en_paras = []
        for i in range(max(0, content_start - 5), len(doc.paragraphs)):
            t = doc.paragraphs[i].text.strip()
            if is_content_en(t):
                en_paras.append(t)

        return en_paras

    return []

# Map parts to volume numbers
# Part 1 covers Orach Chaim vol 1-16 approximately
# We need to figure out the exact mapping
PART_VOLS = {
    'part-1': list(range(1, 17)),
    'part-2': list(range(17, 29)),
    'part-3': list(range(29, 38)),
}

# Collect EN paragraphs per part
part_en = {}
for part, vols in PART_VOLS.items():
    all_en = []
    for v in vols:
        en = extract_en_for_volume(v)
        all_en.extend(en)
        print(f"  Vol {v}: {len(en)} EN content paras")
    part_en[part] = all_en
    print(f"  {part} total: {len(all_en)} EN paras")

# Now process JSON and fix
print("\n=== Fixing JSON ===")

for part_dir in sorted(os.listdir(READER_DIR)):
    if not part_dir.startswith('part-'): continue
    part_path = os.path.join(READER_DIR, part_dir)

    # Get all content segments
    jsfiles = sorted([f for f in os.listdir(part_path)
                      if f.endswith('.json') and f != 'index.json'])

    # Build list of (file, seg_index, segment) for content only
    content_items = []
    for jf in jsfiles:
        data = json.load(open(os.path.join(part_path, jf)))
        for i, seg in enumerate(data['segments']):
            he = seg.get('he', '').strip()
            en = seg.get('en', '').strip()
            if he and len(he) > 15 and not has_hebrew(he):
                # Content segment (non-Hebrew header)
                content_items.append((jf, i, seg))

    # Get docx EN for this part
    en_list = part_en.get(part_dir, [])

    print(f"\n{part_dir}: {len(content_items)} content segs, {len(en_list)} docx EN paras")

    if len(en_list) == 0:
        # Try all volumes
        en_list = []
        for vol_list in PART_VOLS.values():
            for v in vol_list:
                en_list.extend(extract_en_for_volume(v))

    # Fix by position
    fixed = 0
    for idx, (jf, seg_idx, seg) in enumerate(content_items):
        if idx >= len(en_list):
            break
        old_en = seg.get('en', '')
        new_en = en_list[idx]
        if new_en != old_en and len(new_en) > 15:
            seg['en'] = new_en
            fixed += 1

            # Write back to file
            filepath = os.path.join(part_path, jf)
            data = json.load(open(filepath))
            data['segments'][seg_idx] = seg
            json.dump(data, open(filepath, 'w'), indent=2, ensure_ascii=False)

    print(f"  Fixed: {fixed}")

print("\nDone!")