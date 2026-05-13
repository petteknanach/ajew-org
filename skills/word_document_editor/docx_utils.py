"""
Utility functions for python-docx operations.
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def create_document(title="Document", author=""):
    """Create a new Word document with optional title and author."""
    doc = Document()
    if title:
        doc.add_heading(title, 0)
    if author:
        doc.core_properties.author = author
    return doc


def add_heading(doc, text, level=1):
    """Add a heading (1-9) to the document."""
    doc.add_heading(text, level=level)


def add_paragraph(doc, text, bold=False, italic=False, font_size=None, alignment=None):
    """Add a paragraph with optional formatting."""
    p = doc.add_paragraph(text)
    run = p.runs[0] if p.runs else None
    if run:
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if font_size:
            run.font.size = Pt(font_size)
    if alignment == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return doc


def add_list(doc, items, list_type="bullet"):
    """Add a bulleted or numbered list."""
    if list_type == "bullet":
        for item in items:
            doc.add_paragraph(item, style='ListBullet')
    else:
        for item in items:
            doc.add_paragraph(item, style='ListNumber')
    return doc


def add_table(doc, rows, cols, data=None):
    """Add a table with optional cell data."""
    table = doc.add_table(rows=rows, cols=cols)
    if data:
        for i, row_data in enumerate(data):
            for j, val in enumerate(row_data):
                if i < rows and j < cols:
                    table.cell(i, j).text = str(val)
    return doc


def add_image(doc, path, width_inches=None):
    """Add an image to the document."""
    if not os.path.exists(path):
        raise FileNotFoundError(f"Image not found: {path}")
    width = Inches(width_inches) if width_inches else None
    doc.add_picture(path, width=width)
    return doc


def save_document(doc, path):
    """Save the document to a .docx file."""
    doc.save(path)
    return path


def load_document(path):
    """Load an existing .docx document."""
    if not os.path.exists(path):
        raise FileNotFoundError(f"Document not found: {path}")
    return Document(path)
