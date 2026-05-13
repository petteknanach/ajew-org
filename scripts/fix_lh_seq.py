#!/usr/bin/env python3
"""
Fix LH EN-HE pairing using sequential matching.

KEY INSIGHT: Both the JSON and the docx have segments in the SAME ORDER.
The docx alternates [HE, EN, HE, EN, ...] after headers.
The JSON has alternating [HE, EN, HE, EN, ...] segments.

The problem is that the docx has some extra paragraphs (headers, repeated sections)
that cause misalignment. We need to:
1. Filter docx to only content paragraphs
2. Map them 1:1 with JSON segments
"""
from docx import Document
import json
import os

DOCX_DIR = '/mnt/c/Users/Pettek/Documents/Claude Desktop projects/Finished/temp folder/Likutay Halachos'

def is_hebrew_char(c): return '\u05D0' <= c <= '\u05EA'
def has_hebrew(t): return any(is_hebrew_char(c) for c in t)

# Header/metadata patterns
SKIP_PREFIXES = ['hilchos ', 'na nach', 'naanach', 'siman ', 'seif ', 'osio ',
                 'sof ', 'torah ', 'likutay', 'volume ', 'introduction',
                 'a collection', 'the laws ', 'oc ', 'yd ', 'eh ', 'cm ',
                 'orach chaim', 'yoreh deah', 'even haezer', 'choshen mishpat',
                 'hh', 'by our master', 'naanaa', 'nnmmm', 'segment',
                 'one stop', 'arranged by', 'jaw', 'for any', 'all corrections']

def is_paragraph_header(text):
    t = text.lower().strip()
    if len(t) < 5: return True
    words = t.split()
    if len(words) <= 2: return True
    return any(t.startswith(p) for p in SKIP_PREFIXES)

def extract_docx_as_sections(docx_path):
    """Extract docx as alternating (HE, EN) content sections."""
    doc = Document(docx_path)
    paras = [p.text for p in doc.paragraphs]

    # First pass: classify each paragraph
    sections = []
    for p in paras:
        p = p.strip()
        if len(p) < 10:
            sections.append(('empty', p))
        elif is_paragraph_header(p):
            sections.append(('header', p))
        elif has_hebrew(p):
            sections.append(('he', p))
        else:
            sections.append(('en', p))

    return sections

def parse_docx_into_pairs(docx_path):
    """Parse docx into (he_text, en_text) content pairs."""
    sections = extract_docx_as_sections(docx_path)
    pairs = []
    current_he = None
    current_en_parts = []
    found_first_he = False

    for typ, text in sections:
        if typ == 'header':
            if current_he is not None:
                # End current pair
                if current_en_parts:
                    pairs.append((current_he, '\n'.join(current_en_parts)))
                    current_he = None
                    current_en_parts = []
                else:
                    current_he = None
            continue

        if typ == 'he':
            if current_he is not None and found_first_he:
                # Save previous pair if it has EN
                if current_en_parts:
                    pairs.append((current_he, '\n'.join(current_en_parts)))
                elif current_he:
                    pairs.append((current_he, ''))
            current_he = text
            current_en_parts = []
            found_first_he = True

        elif typ == 'en' and found_first_he:
            current_en_parts.append(text)

        elif typ == 'empty' and found_first_he:
            current_en_parts.append(text)

    # Don't forget last pair
    if current_he and found_first_he:
        pairs.append((current_he, '\n'.join(current_en_parts) if current_en_parts else ''))

    return pairs

def norm_text(t):
    import re
    t = re.sub(r'[\u0591-\u05BD\u05BF-\u05C7]', '', t)  # strip nikkud
    return re.sub(r'\s+', ' ', t.lower()).strip()

def parse_docx_into_pairs(docx_path):
    """Parse docx into (he_text, en_text) content pairs."""
    doc = Document(docx_path)
    paras = [p.text for p in doc.paragraphs]

    def is_hebrew_char(c): return '\u05D0' <= c <= '\u05EA'
    def has_hebrew(t): return any(is_hebrew_char(c) for c in t)

    SKIP_PREFIXES = ['hilchos ','na nach','siman ','seif ','osio ','volume ',
                     'introduction','likutay','a collection','the laws ']

    def is_header(t):
        tl = t.lower().strip()
        if len(tl) < 5 or len(tl.split()) <= 2: return True
        return any(tl.startswith(p) for p in SKIP_PREFIXES)

    sections = []
    for p in paras:
        p = p.strip()
        if len(p) < 10:
            sections.append(('empty', p))
        elif is_header(p):
            sections.append(('header', p))
        elif has_hebrew(p):
            sections.append(('he', p))
        else:
            sections.append(('en', p))

    pairs = []
    current_he = None
    current_en_parts = []
    found_first_he = False

    for typ, text in sections:
        if typ == 'header':
            if current_he and current_en_parts:
                pairs.append((current_he, '\n'.join(current_en_parts)))
            current_he = None
            current_en_parts = []
            continue
        if typ == 'he':
            if current_he and found_first_he and current_en_parts:
                pairs.append((current_he, '\n'.join(current_en_parts)))
            current_he = text
            current_en_parts = []
            found_first_he = True
        elif typ == 'en' and found_first_he:
            current_en_parts.append(text)
        elif typ == 'empty' and found_first_he:
            current_en_parts.append(text)

    if current_he and found_first_he:
        pairs.append((current_he, '\n'.join(current_en_parts) if current_en_parts else ''))

    return pairs

def fix_lh():
    import json
    print("Parsing docx files...")
    all_pairs = []
    for df in sorted(os.listdir(DOCX_DIR)):
        if not df.endswith('.docx'): continue
        pairs = parse_docx_into_pairs(os.path.join(DOCX_DIR, df))
        all_pairs.extend(pairs)
    print(f"  Total docx content pairs: {len(all_pairs)}")

    # Filter to pairs that have both HE and EN
    valid_pairs = [(h, e) for h, e in all_pairs if h and e]
    print(f"  Valid (HE+EN) pairs: {len(valid_pairs)}")

    # Flatten: create a sequential list of EN paragraphs
    en_seq = [e for h, e in valid_pairs]

    for part_dir in ['part-1', 'part-2', 'part-3', 'part-4', 'part-5', 'part-6', 'part-7', 'part-8']:
        part_path = os.path.join('/root/ajew-org/public/reader/likutay-halachos', part_dir)
        if not os.path.exists(part_path): continue

        jsfiles = sorted([f for f in os.listdir(part_path)
                          if f.endswith('.json') and f != 'index.json'])

        # Collect all segment EN texts in order
        all_seg_ens = []
        for jf in jsfiles:
            data = json.load(open(os.path.join(part_path, jf)))
            for seg in data['segments']:
                if seg.get('he','').strip() and not is_header(seg.get('he','').strip()):
                    all_seg_ens.append(seg.get('en','').strip())

        print(f"\n  {part_dir}: {len(all_seg_ens)} content segments")

        # Compare seq lengths
        print(f"    Docx EN count: {len(en_seq)}")
        print(f"    JSON EN count: {len(all_seg_ens)}")

        # Check: do they differ?
        if len(all_seg_ens) == len(en_seq):
            # Perfect match - assign directly
            idx = 0
            for jf in jsfiles:
                filepath = os.path.join(part_path, jf)
                data = json.load(open(filepath))
                changed = False
                for seg in data['segments']:
                    he = seg.get('he','').strip()
                    if he and not is_header(he):
                        if idx < len(en_seq):
                            seg['en'] = en_seq[idx]
                            changed = True
                        idx += 1
                if changed:
                    json.dump(data, open(filepath, 'w'), indent=2, ensure_ascii=False)
            print(f"    Fixed by direct sequential assignment")
        else:
            print(f"    Length mismatch - need word-level matching")

fix_lh()