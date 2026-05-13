#!/usr/bin/env python3
"""Explore LH docx structure in detail - all paragraphs."""
from docx import Document

docx_path = '/mnt/c/Users/Pettek/Documents/Claude Desktop projects/Finished/temp folder/Likutay Halachos/Volume_01_OC1_English.docx'
doc = Document(docx_path)

print("ALL paragraphs in LH docx:")
print("=" * 100)
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    has_he = any(ord(c) > 127 for c in text)
    style = p.style.name if p.style else 'Normal'
    marker = "# HE " if has_he else "    EN"
    print(f"[{i:3d}] {marker} | {text[:200]}")

print()
print(f"Total paragraphs: {len(doc.paragraphs)}")