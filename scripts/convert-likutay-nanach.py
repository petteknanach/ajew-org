#!/usr/bin/env python3
"""
Convert Likutay Nanach DOCX files to text and HTML for ajew.org
"""
import os
import json
import re
from pathlib import Path
from docx import Document

def extract_docx_text(docx_path):
    """Extract text from DOCX file"""
    try:
        doc = Document(docx_path)
        full_text = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n\n".join(full_text)
    except Exception as e:
        print(f"Error processing {docx_path}: {e}")
        return ""

def detect_chapters(text, volume_num):
    """Detect chapters/sections in the text"""
    chapters = []
    
    # Common Hebrew chapter markers
    chapter_patterns = [
        r'^סימן\s+(\S+)',  # Siman
        r'^פרק\s+(\S+)',   # Perek
        r'^הלכה\s+(\S+)',  # Halacha
        r'^אות\s+(\S+)',   # Ot
        r'^(\d+)\.',       # Number with dot
        r'^\[(\d+)\]',     # Number in brackets
    ]
    
    lines = text.split('\n')
    current_chapter = None
    current_content = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if this line starts a new chapter
        is_chapter_start = False
        chapter_title = None
        
        for pattern in chapter_patterns:
            match = re.match(pattern, line)
            if match:
                is_chapter_start = True
                chapter_num = match.group(1)
                chapter_title = f"Chapter {chapter_num}"
                break
        
        if is_chapter_start and current_chapter is not None:
            # Save previous chapter
            if current_content:
                chapters.append({
                    'title': current_chapter['title'],
                    'number': current_chapter['number'],
                    'content': '\n'.join(current_content)
                })
            current_content = [line]
            current_chapter = {
                'title': chapter_title,
                'number': chapter_num
            }
        elif is_chapter_start:
            # First chapter
            current_chapter = {
                'title': chapter_title,
                'number': chapter_num
            }
            current_content = [line]
        elif current_chapter is not None:
            # Continue current chapter
            current_content.append(line)
        else:
            # Content before first chapter
            if not chapters:
                chapters.append({
                    'title': 'Introduction',
                    'number': '0',
                    'content': line
                })
            else:
                chapters[0]['content'] += '\n' + line
    
    # Add last chapter
    if current_chapter and current_content:
        chapters.append({
            'title': current_chapter['title'],
            'number': current_chapter['number'],
            'content': '\n'.join(current_content)
        })
    
    # If no chapters detected, create one big chapter
    if not chapters:
        chapters.append({
            'title': f'Volume {volume_num}',
            'number': '1',
            'content': text
        })
    
    return chapters

def create_html_chapter(chapter, chapter_num, total_chapters, volume_num):
    """Create HTML for a chapter"""
    title = chapter['title']
    content = chapter['content']
    
    # Basic HTML structure
    html = f"""<!DOCTYPE html>
<html lang="he" dir="rtl">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>לקוטי ננח - כרך {volume_num} - {title}</title>
    <style>
        body {{
            font-family: 'Arial', 'David', sans-serif;
            line-height: 1.8;
            margin: 0;
            padding: 20px;
            background: #f8f9fa;
            color: #333;
        }}
        .container {{
            max-width: 800px;
            margin: 0 auto;
            background: white;
            padding: 30px;
            border-radius: 10px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
            text-align: center;
        }}
        .content {{
            font-size: 18px;
            text-align: justify;
            white-space: pre-line;
        }}
        .nav {{
            display: flex;
            justify-content: space-between;
            margin-top: 30px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
        }}
        .nav a {{
            background: #3498db;
            color: white;
            padding: 10px 20px;
            text-decoration: none;
            border-radius: 5px;
            font-weight: bold;
        }}
        .nav a:hover {{
            background: #2980b9;
        }}
        .chapter-info {{
            text-align: center;
            color: #7f8c8d;
            margin-bottom: 20px;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>לקוטי ננח - כרך {volume_num}</h1>
        <div class="chapter-info">
            <h2>{title}</h2>
            <p>פרק {chapter_num} מתוך {total_chapters}</p>
        </div>
        
        <div class="content">
{content}
        </div>
        
        <div class="nav">
            {f'<a href="chapter{int(chapter_num)-1}.html">← הפרק הקודם</a>' if int(chapter_num) > 1 else '<span></span>'}
            <a href="index.html">📚 חזרה לתוכן העניינים</a>
            {f'<a href="chapter{int(chapter_num)+1}.html">הפרק הבא →</a>' if int(chapter_num) < total_chapters else '<span></span>'}
        </div>
    </div>
</body>
</html>"""
    
    return html

def create_index_html(volume, chapters):
    """Create index page for a volume"""
    volume_num = volume['number']
    volume_name = volume['name']
    
    # Create chapter links
    chapter_links = ""
    for i, chapter in enumerate(chapters, 1):
        chapter_links += f'        <li><a href="chapter{i}.html">{chapter["title"]}</a></li>\n'
    
    html = f"""<!DOCTYPE html>
<html lang="he" dir="rtl">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>לקוטי ננח - כרך {volume_num}</title>
    <style>
        body {{
            font-family: 'Arial', 'David', sans-serif;
            line-height: 1.8;
            margin: 0;
            padding: 20px;
            background: #f8f9fa;
            color: #333;
        }}
        .container {{
            max-width: 800px;
            margin: 0 auto;
            background: white;
            padding: 30px;
            border-radius: 10px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
            text-align: center;
        }}
        .subtitle {{
            text-align: center;
            color: #7f8c8d;
            font-size: 1.2em;
            margin-bottom: 30px;
        }}
        .toc {{
            background: #f8f9fa;
            padding: 20px;
            border-radius: 5px;
            margin: 20px 0;
        }}
        .toc h2 {{
            color: #2c3e50;
            margin-top: 0;
        }}
        .toc ul {{
            list-style: none;
            padding: 0;
        }}
        .toc li {{
            margin: 10px 0;
            padding: 10px;
            background: white;
            border-radius: 5px;
            border-left: 4px solid #3498db;
        }}
        .toc a {{
            color: #2c3e50;
            text-decoration: none;
            font-weight: bold;
            display: block;
        }}
        .toc a:hover {{
            color: #3498db;
        }}
        .back-link {{
            display: inline-block;
            margin-top: 20px;
            background: #2c3e50;
            color: white;
            padding: 10px 20px;
            text-decoration: none;
            border-radius: 5px;
        }}
        .back-link:hover {{
            background: #34495e;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>לקוטי ננח</h1>
        <div class="subtitle">
            <h2>{volume_name}</h2>
            <p>כרך {volume_num} - {len(chapters)} פרקים</p>
        </div>
        
        <div class="intro">
            <p><strong>לקוטי ננח</strong> הוא אוסף של תורות, שיחות והנהגות מגדולי חסידי ברסלב, בעיקר סביב עניין ה"ננח" - נא נח נחמ נחמן מאומן.</p>
            <p>הספר כולל דברי תורה, סיפורים, והנהגות הקשורות לדרך המיוחדת של נא נח נחמ נחמן מאומן.</p>
        </div>
        
        <div class="toc">
            <h2>תוכן העניינים</h2>
            <ul>
{chapter_links}
            </ul>
        </div>
        
        <a href="/books" class="back-link">← חזרה לכל הספרים</a>
    </div>
</body>
</html>"""
    
    return html

def main():
    # Paths
    docx_path = Path("C:/Users/Pettek/Documents/Likutay Nanach")
    output_path = Path("C:/Users/Pettek/.openclaw/workspace/ajew-org/public/books/likutay-nanach")
    
    # Create output directory
    output_path.mkdir(parents=True, exist_ok=True)
    
    # Find all DOCX files
    docx_files = list(docx_path.glob("*.docx"))
    print(f"Found {len(docx_files)} Likutay Nanach volumes:")
    
    volumes_data = []
    
    for docx_file in docx_files:
        print(f"\nProcessing: {docx_file.name}")
        
        # Extract volume number from filename
        vol_match = re.search(r'vol\s*(\d+)', docx_file.name, re.IGNORECASE)
        volume_num = vol_match.group(1) if vol_match else "unknown"
        
        volume_name = f"Likutay Nanach Volume {volume_num}"
        
        # Extract text
        text = extract_docx_text(docx_file)
        print(f"  Extracted {len(text)} characters")
        
        if not text:
            print(f"  Warning: No text extracted from {docx_file.name}")
            continue
        
        # Detect chapters
        chapters = detect_chapters(text, volume_num)
        print(f"  Detected {len(chapters)} chapters")
        
        # Create volume directory
        volume_dir = output_path / f"volume-{volume_num}"
        volume_dir.mkdir(exist_ok=True)
        
        # Save raw text
        raw_text_file = volume_dir / "raw-text.txt"
        raw_text_file.write_text(text, encoding='utf-8')
        
        # Create HTML pages for each chapter
        for i, chapter in enumerate(chapters, 1):
            html = create_html_chapter(chapter, i, len(chapters), volume_num)
            chapter_file = volume_dir / f"chapter{i}.html"
            chapter_file.write_text(html, encoding='utf-8')
        
        # Create index page
        index_html = create_index_html({
            'number': volume_num,
            'name': volume_name
        }, chapters)
        
        index_file = volume_dir / "index.html"
        index_file.write_text(index_html, encoding='utf-8')
        
        # Save chapter metadata
        metadata = {
            'volume': volume_num,
            'name': volume_name,
            'source_file': docx_file.name,
            'chapters': chapters,
            'total_chapters': len(chapters)
        }
        
        metadata_file = volume_dir / "metadata.json"
        metadata_file.write_text(json.dumps(metadata, ensure_ascii=False, indent=2), encoding='utf-8')
        
        volumes_data.append({
            'volume': volume_num,
            'name': volume_name,
            'path': f"/books/likutay-nanach/volume-{volume_num}/",
            'chapters': len(chapters)
        })
        
        print(f"  Created {len(chapters)} chapter pages in {volume_dir}")
    
    # Create master index
    master_index = {
        'volumes': volumes_data,
        'total_volumes': len(volumes_data),
        'total_chapters': sum(v['chapters'] for v in volumes_data)
    }
    
    master_index_file = output_path / "index.json"
    master_index_file.write_text(json.dumps(master_index, ensure_ascii=False, indent=2), encoding='utf-8')
    
    print(f"\n{'='*60}")
    print(f"Conversion complete!")
    print(f"Total volumes: {len(volumes_data)}")
    print(f"Total chapters: {sum(v['chapters'] for v in volumes_data)}")
    print(f"Output directory: {output_path}")
    print(f"{'='*60}")
    
    # Create integration instructions for ajew.org
    print("\nNext steps for ajew.org integration:")
    print("1. Create Astro pages in src/pages/books/likutay-nanach/")
    print("2. Link to HTML files from public/books/likutay-nanach/")
    print("3. Add to site search index")
    print("4. Update books navigation")

if __name__ == "__main__":
    main()