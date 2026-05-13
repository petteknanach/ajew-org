#!/usr/bin/env python3
"""Analyze docx Volume 1 structure in detail."""
from docx import Document
import os

DOCX_DIR = '/mnt/c/Users/Pettek/Documents/Claude Desktop projects/Finished/temp folder/Likutay Halachos'

doc = Document(os.path.join(DOCX_DIR, 'Volume_01_OC1_English.docx'))
print(f"Volume 01 - {len(doc.paragraphs)} paragraphs\n")

# Show ALL paragraphs
for i, p in enumerate(doc.paragraphs[:80]):
    t = p.text.strip()
    if not t: continue
    print(f"P{i:3d} | {len(t):4d} | {t[:100]}")