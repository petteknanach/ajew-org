#!/usr/bin/env python3
"""
Fix LH pairing v13 - Proper structural alignment.

Key insight: JSON segments correspond to halachot/torah content.
Docx has front matter + content. Need to:
1. Find where actual content starts in docx (after front matter)
2. Match docx EN paragraphs to JSON segments AFTER the front matter
"""
from docx import Document
import json, os, re

DOCX_DIR = '/mnt/c/Users/Pettek/Documents/Claude Desktop projects/Finished/temp folder/Likutay Halachos'
READER_DIR = '/root/ajew-org/public/reader/likutay-halachos'

def is_hebrew_char(c): return '\u05D0' <= c <= '\u05EA'
def has_hebrew(t): return any(is_hebrew_char(c) for c in t)
def strip_nikkud(t): return re.sub(r'[\u0591-\u05BD\u05BF-\u05C7]', '', t)
def norm(t): return re.sub(r'\s+', ' ', strip_nikkud(t.lower())).strip()
def he_words(t): return set(re.findall(r'[\u05D0-\u05EA]{4,}', norm(t)))

# Patterns that mark the BEGINNING of actual content (not front matter)
CONTENT_START_MARKERS = [
    'סוף',  # Sof - end
    'שם',   # Shem - name/reference
]

def is_front_matcher(text):
    """Check if this is front matter (not actual content)."""
    t = text.lower().strip()
    # Short text, likely header
    if len(t.split()) <= 2: return True
    # Known front matter patterns
    front_patterns = [
        'hilchos ', 'na nach', 'siman ', 'seif ', 'osio ',
        'volume ', 'introduction', 'likutay', 'a collection',
        'the laws ', 'oc ', 'yd ', 'eh ', 'cm ', 'like all',
        'naanach', 'segment', 'one stop', 'arranged by',
        'nitan', 'rabbi nachman', 'the entire likutay',
        'each paragraph in this volume',
        'table of contents', 'right-click',
        'copyright', 'rough draft', 'no copyright',
        'free for all', 'books of rabbi nachman',
        'character — the aleph', 'the stories of rabbi',
        'outpouring of the soul', 'the fires of israel',
        'rabbi nachman of breslov: who', 'live up the good',
        'pray with your limbs', 'the praises of rabbi',
        'likutay halachos (complete', 'ajew.org',
        'one stop for everything',
        'bas sheva rosa', 'bas sheva', 'yisroel dov',
        'student of rabbi', 'na naach', 'na nach nachma',
        'note on paragraph numbers',
        'cross-reference number',
    ]
    return any(t.startswith(p) or p in t for p in front_patterns)

def find_content_start_index(doc):
    """Find the paragraph index where actual content begins."""
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if len(t) < 15: continue

        # The content typically starts with a Hebrew text beginning
        # with "והנה" or "הפעם" or similar
        if has_hebrew(t):
            hw = set(re.findall(r'[\u05D0-\u05EA]{4,}', strip_nikkud(t.lower())))
            # Check if this looks like content (has distinctive Hebrew words)
            if hw and not is_front_matcher(t):
                return i
    return 0

def extract_docx_content_paras(doc_path):
    """Extract docx paragraphs after front matter."""
    doc = Document(doc_path)

    start_idx = find_content_start_index(doc)
    print(f"    Content starts at paragraph {start_idx}")

    content_paras = []
    for i in range(start_idx, len(doc.paragraphs)):
        t = doc.paragraphs[i].text.strip()
        if len(t) < 5: continue

        if has_hebrew(t):
            if not is_front_matcher(t):
                content_paras.append(('he', t))
        else:
            # English paragraph - but skip if it looks like metadata
            if len(t) > 15 and not is_front_matcher(t):
                content_paras.append(('en', t))

    return content_paras

def main():
    print("=== Building docx content index ===\n")

    all_content = []
    for df in sorted(os.listdir(DOCX_DIR)):
        if not df.endswith('.docx'): continue
        path = os.path.join(DOCX_DIR, df)
        paras = extract_docx_content_paras(path)
        all_content.extend(paras)
        he_count = sum(1 for t, _ in paras if t == 'he')
        en_count = sum(1 for t, _ in paras if t == 'en')
        print(f"  {df}: {he_count} HE, {en_count} EN content paragraphs")

    print(f"\nTotal: {len(all_content)} content paragraphs")

    # Separate HE and EN
    he_paras = [t for t, _ in all_content if t == 'he']
    en_paras = [t for t, _ in all_content if t == 'en']

    print(f"  HE paragraphs: {len(he_paras)}")
    print(f"  EN paragraphs: {len(en_paras)}")

    # Now process JSON
    print("\n=== Processing JSON segments ===\n")

    json_he_count = 0
    for part_dir in sorted(os.listdir(READER_DIR)):
        if not part_dir.startswith('part-'): continue

        part_path = os.path.join(READER_DIR, part_dir)
        jsfiles = sorted([f for f in os.listdir(part_path)
                          if f.endswith('.json') and f != 'index.json'])

        for jf in jsfiles:
            data = json.load(open(os.path.join(part_path, jf)))
            for seg in data['segments']:
                he = seg.get('he', '').strip()
                if he and len(he) > 20 and not is_front_matcher(he):
                    json_he_count += 1

    print(f"Total JSON content segments: {json_he_count}")
    print(f"Docx EN paragraphs: {len(en_paras)}")
    print(f"\nIf these match, we can assign EN by position")

    # For now, just show the numbers
    print(f"\n=== Matching Analysis ===")
    print(f"JSON has {json_he_count} content segments")
    print(f"Docx has {len(en_paras)} EN content paragraphs")
    print(f"Difference: {abs(json_he_count - len(en_paras))}")

main()